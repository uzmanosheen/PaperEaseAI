import re

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

URDU_FONT = "Noto Nastaliq Urdu"   
ENGLISH_FONT = "Times New Roman"

HEADER_LABEL_PATTERNS = [
    r"school\s*name\s*:", r"subject\s*:", r"class\s*:", r"term\s*:",
    r"session\s*:", r"total\s*marks\s*:", r"time\s*allowed\s*:", r"time\s*:",
    r"roll\s*(no|number)\s*:", r"student'?s?\s*name\s*:", r"name\s*:", r"date\s*:",
]

# Matches strings like "a) oxygen", "B) 9", " c)  none "
_MCQ_OPTION_RE = re.compile(r"^\s*([a-e])\)\s*(.*)$", re.IGNORECASE)


def strip_header_lines(text):
    """Remove any leftover masthead-style lines (e.g. 'Roll No: ___') from
    extracted text, since that info is rendered separately in the header table."""
    if not text:
        return text
    lines = text.split("\n")
    kept = [
        line for line in lines
        if not any(re.search(p, line, re.IGNORECASE) for p in HEADER_LABEL_PATTERNS)
    ]
    return "\n".join(kept).strip()


def sanitize_data(data):
    data = dict(data)
    data["exam_title"] = strip_header_lines(data.get("exam_title", ""))
    data["instructions"] = [
        line for line in (strip_header_lines(i) for i in data.get("instructions", []))
        if line
    ]
    for sec in data.get("sections", []):
        cleaned = strip_header_lines(sec.get("section_title", ""))
        if cleaned:
            sec["section_title"] = cleaned
    return data


def _looks_like_mcq_options(sub_parts):
    """Return True when every sub_part looks like an MCQ option (a) ... b) ...)."""
    if not sub_parts or len(sub_parts) < 2:
        return False
    return all(_MCQ_OPTION_RE.match(sp) for sp in sub_parts)


def _format_mcq_options(sub_parts, max_single_line_chars=55):
    """Format MCQ options inline, wrapping to two options per line when too long.

    - Short option sets are placed on a single line.
    - Longer sets are split into two options per line.
    - One tab is used between short options; two tabs are used after long options
      to keep the layout visually balanced.
    """
    options = []
    for sp in sub_parts:
        m = _MCQ_OPTION_RE.match(sp)
        letter = m.group(1).lower()
        text = m.group(2).strip()
        options.append((letter, text))

    rendered = [f"{letter}) {text}" for letter, text in options]
    # Rough character budget: if the rendered options fit comfortably on one
    # line, keep them there.
    total_len = sum(len(r) for r in rendered) + (len(rendered) - 1)

    use_single_line = len(options) <= 2 or total_len <= max_single_line_chars

    groups = [rendered] if use_single_line else [
        rendered[i:i + 2] for i in range(0, len(rendered), 2)
    ]

    final_lines = []
    for group in groups:
        line_parts = []
        for i, item in enumerate(group):
            if i < len(group) - 1:
                # Wider spacing after longer options.
                sep = "\t\t" if len(item) > 15 else "\t"
                line_parts.append(item + sep)
            else:
                line_parts.append(item)
        final_lines.append("".join(line_parts))

    return "\n".join(final_lines)


def set_rtl(paragraph):
    """Mark a paragraph as right-to-left (needed for Urdu)."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def set_font(run, font_name, size=12, bold=False, urdu=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if urdu:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:cs"), font_name)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        lang = OxmlElement("w:lang")
        lang.set(qn("w:bidi"), "ur-PK")
        rPr.append(lang)


def shade(paragraph, color="D9D9D9"):
    """Light grey background behind a paragraph (used for Template 2 section headers)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def _set_cell_text(cell, text, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    return p


def add_header_table(doc, header_info, logo_path=None):
    """Default masthead: logo | school name/term title, then a 4-column
    Subject/Class/Name/Roll/Time/Marks/Date grid — matches the standard
    school exam header layout."""
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    logo_cell = table.cell(0, 0)
    if logo_path:
        p = logo_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(0.9))

    title_cell = table.cell(0, 1)
    title_cell.merge(table.cell(0, 3))
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(header_info.get("school_name", ""))
    run.bold = True
    run.font.size = Pt(16)

    subtitle_parts = [p for p in [header_info.get("term", ""), header_info.get("session", "")] if p]
    if subtitle_parts:
        sub_p = title_cell.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = sub_p.add_run(f"({' — '.join(subtitle_parts)})")
        r2.bold = True
        r2.font.size = Pt(12)

    grid = [
        ("Subject:", header_info.get("subject", ""), "Class:", header_info.get("class_name", "")),
        ("Student's Name:", "", "Time Allowed:", header_info.get("time", "")),
        ("Roll Number:", "", "Total Marks:", header_info.get("total_marks", "")),
        ("Date:", header_info.get("date", ""), "Obtained Marks:", ""),
    ]
    for row_idx, (l1, v1, l2, v2) in enumerate(grid, start=1):
        cells = table.rows[row_idx].cells
        _set_cell_text(cells[0], l1, bold=True)
        _set_cell_text(cells[1], v1)
        _set_cell_text(cells[2], l2, bold=True)
        _set_cell_text(cells[3], v2)

    doc.add_paragraph()


def add_body(doc, data, template="template1"):
    """The exam title, instructions, sections and questions. Shared by both
    built-in templates AND the custom-template flow (see custom_template.py)."""
    data = sanitize_data(data)
    is_urdu = data.get("language", "english").lower() == "urdu"
    font = URDU_FONT if is_urdu else ENGLISH_FONT

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(data.get("exam_title", ""))
    set_font(run, font, 15, bold=True, urdu=is_urdu)
    if is_urdu:
        set_rtl(title_p)

    if data.get("instructions"):
        label_p = doc.add_paragraph()
        r = label_p.add_run("ہدایات:" if is_urdu else "Instructions:")
        set_font(r, font, 11, bold=True, urdu=is_urdu)
        if is_urdu:
            set_rtl(label_p)
        for ins in data["instructions"]:
            b = doc.add_paragraph(style="List Bullet")
            r = b.add_run(ins)
            set_font(r, font, 11, urdu=is_urdu)
            if is_urdu:
                set_rtl(b)

    for sec in data.get("sections", []):
        sec_p = doc.add_paragraph()
        sec_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        marks_txt = f"  ({sec.get('section_marks','')})" if sec.get("section_marks") else ""
        r = sec_p.add_run(f"{sec.get('section_title','')}{marks_txt}")
        set_font(r, font, 13, bold=True, urdu=is_urdu)
        if template == "template2":
            shade(sec_p)
        if is_urdu:
            set_rtl(sec_p)

        for q in sec.get("questions", []):
            q_p = doc.add_paragraph()
            marks_str = f"   [{q.get('marks','')}]" if q.get("marks") else ""
            r = q_p.add_run(f"{q.get('number','')}. {q.get('text','')}{marks_str}")
            set_font(r, font, 12, urdu=is_urdu)
            if is_urdu:
                set_rtl(q_p)
            sub_parts = q.get("sub_parts", [])
            if _looks_like_mcq_options(sub_parts):
                mcq_text = _format_mcq_options(sub_parts)
                opt_p = doc.add_paragraph()
                opt_p.paragraph_format.left_indent = Cm(1)
                r2 = opt_p.add_run(mcq_text)
                set_font(r2, font, 11, urdu=is_urdu)
                if is_urdu:
                    set_rtl(opt_p)
                # Column tab stops so options line up neatly.
                tabs = opt_p.paragraph_format.tab_stops
                tabs.add_tab_stop(Cm(4.5))
                tabs.add_tab_stop(Cm(9))
                tabs.add_tab_stop(Cm(13.5))
            else:
                for sp in sub_parts:
                    sp_p = doc.add_paragraph()
                    sp_p.paragraph_format.left_indent = Cm(1)
                    r2 = sp_p.add_run(sp)
                    set_font(r2, font, 11, urdu=is_urdu)
                    if is_urdu:
                        set_rtl(sp_p)


def build_docx(data, header_info, logo_path, template="template1", output_path="output.docx"):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    add_header_table(doc, header_info, logo_path)
    add_body(doc, data, template)

    doc.save(output_path)
    return output_path